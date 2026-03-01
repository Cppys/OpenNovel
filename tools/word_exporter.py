"""Export novel or short story to Word (.docx) format."""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.database import Database

logger = logging.getLogger(__name__)


def _add_cover_page(doc: Document, title: str, genre: str, synopsis: str):
    """Add a simple cover/title page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.bold = True

    # Genre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(genre)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Synopsis
    if synopsis:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(synopsis)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def export_novel_to_word(
    db: Database,
    novel_id: int,
    output_path: Optional[Path] = None,
) -> Path:
    """Export a novel with all chapters to a Word document.

    Args:
        db: Database instance.
        novel_id: The novel ID to export.
        output_path: Where to save the .docx. Auto-generated if None.

    Returns:
        Path to the saved .docx file.

    Raises:
        ValueError: If novel not found or has no chapters.
    """
    novel = db.get_novel(novel_id)
    if not novel:
        raise ValueError(f"未找到ID为 {novel_id} 的小说")

    chapters = db.get_chapters(novel_id)
    if not chapters:
        raise ValueError(f"小说《{novel.title}》没有章节内容")

    volumes = db.get_volumes(novel_id)
    vol_map = {v.id: v for v in volumes}

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Pt(24)

    # Cover page
    _add_cover_page(doc, novel.title, novel.genre, novel.synopsis)

    # Group chapters by volume
    current_vol_id = None
    for ch in chapters:
        # Volume heading if volume changed
        if ch.volume_id and ch.volume_id != current_vol_id:
            current_vol_id = ch.volume_id
            vol = vol_map.get(ch.volume_id)
            if vol:
                h = doc.add_heading(vol.title, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Chapter heading
        ch_title = f"第{ch.chapter_number}章"
        if ch.title:
            ch_title += f" {ch.title}"
        h = doc.add_heading(ch_title, level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Chapter content — split by paragraphs
        content = ch.content or ""
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

        # Page break between chapters (except last)
        if ch != chapters[-1]:
            doc.add_page_break()

    # Determine output path
    if output_path is None:
        exports_dir = Path("./exports")
        exports_dir.mkdir(parents=True, exist_ok=True)
        safe_title = novel.title.replace(" ", "_")[:30]
        output_path = exports_dir / f"{safe_title}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Novel exported to %s", output_path)
    return output_path


def export_short_story_to_word(
    db: Database,
    story_id: int,
    output_path: Optional[Path] = None,
) -> Path:
    """Export a short story to a Word document.

    Args:
        db: Database instance.
        story_id: The short story ID to export.
        output_path: Where to save the .docx. Auto-generated if None.

    Returns:
        Path to the saved .docx file.

    Raises:
        ValueError: If story not found or has no content.
    """
    story = db.get_short_story(story_id)
    if not story:
        raise ValueError(f"未找到ID为 {story_id} 的短故事")

    content = story.get("content", "")
    if not content:
        raise ValueError(f"短故事《{story.get('title', '?')}》没有内容")

    title = story.get("title", "未命名短故事")
    genre = story.get("genre", "")
    synopsis = story.get("synopsis", "")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Pt(24)

    # Cover page
    _add_cover_page(doc, title, genre, synopsis)

    # Title heading
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content
    for para_text in content.split("\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    # Determine output path
    if output_path is None:
        exports_dir = Path("./exports")
        exports_dir.mkdir(parents=True, exist_ok=True)
        safe_title = title.replace(" ", "_")[:30]
        output_path = exports_dir / f"{safe_title}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Short story exported to %s", output_path)
    return output_path
